import os

import comtypes.client
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

from utilities import Utils
from jinja2 import Environment, FileSystemLoader
import asyncio
import os
import base64
from pyppeteer import launch
import pyppeteer


class PdfReporting:
    def __init__(self, logo_path, encrypted_template_file_path, data, tc_id, document_name):
        self.utils = Utils()
        self.tc_id = tc_id
        self.logo_path = logo_path
        self.base64_logo = self._encode_logo()
        self.html_content = self._generate_html(encrypted_template_file_path, data)
        self.head_template = self._create_header_template()
        self.footer_template = self._create_footer_template()
        self.document_name_pdf = self.utils.get_test_result_folder() + "\\" + document_name + ".pdf"
    
    def _encode_logo(self):
        with open(self.logo_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def _create_header_template(self):
        return f'''
        <div style="display: flex; width: 100%; height: auto; justify-content: space-between; align-items: center; border-bottom: 2px solid #413a97;">
            <div style="flex: 0;margin-left: 5px;">
                <img id="logo" src="data:image/png;base64,{self.base64_logo}" alt="Logo" style="max-height: 60px;object-fit: contain;">
            </div>
            <div id="tcid" style="flex: 0;font-size:10px;white-space: nowrap;margin-right: 20px;">
                <h1>{self.tc_id}</h1>
            </div>
        </div>
        '''

    def _create_footer_template(self):
        return '''
        <div style='font-size:10px; width:100%; text-align:center;'>
            Page <span class="pageNumber"></span> of <span class="totalPages"></span>
        </div>
        '''

    def _generate_html(self, encrypted_template_file_path, data):
        # Decrypt the file
        decrypted_template = self.utils.decrypt_file(encrypted_template_file_path)

        # Setup Jinja2 environment and render the HTML template
        env = Environment(loader=FileSystemLoader("."))
        template = env.from_string(decrypted_template)
        output = template.render(data)

        # Save the populated HTML to a file
        with open("output.html", "w") as f:
            f.write(output)

        # print("HTML file generated successfully!")
        return output
    
    async def generate_pdf(self):
        browser = await launch({"headless": True})
        # print(await browser.version())
        page = await browser.newPage()

        await page.setContent(self.html_content)

        await page.pdf({
            "path": self.document_name_pdf,
            "format": "Letter",
            "printBackground": True,
            "landscape": True,
            "margin": {
                "top": 85,
                "bottom": 25,
                "left": 25,
                "right": 25
            },
            "displayHeaderFooter": True,
            "headerTemplate": self.head_template,
            "footerTemplate": self.footer_template,
        })

        await browser.close()
    
    # def create_element(self, name):
    #     self.is_not_used()
    #     return OxmlElement(name)

    # def create_attribute(self, element, name, value):
    #     self.is_not_used()
    #     element.set(ns.qn(name), value)

    # def add_page_number(self, paragraph):
    #     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #     page_run = paragraph.add_run()
    #     t1 = self.create_element('w:t')
    #     self.create_attribute(t1, 'xml:space', 'preserve')
    #     t1.text = 'Page '
    #     page_run._r.append(t1)

    #     page_num_run = paragraph.add_run()

    #     fld_char1 = self.create_element('w:fldChar')
    #     self.create_attribute(fld_char1, 'w:fldCharType', 'begin')

    #     instr_text = self.create_element('w:instrText')
    #     self.create_attribute(instr_text, 'xml:space', 'preserve')
    #     instr_text.text = "PAGE"

    #     fld_char2 = self.create_element('w:fldChar')
    #     self.create_attribute(fld_char2, 'w:fldCharType', 'end')

    #     page_num_run._r.append(fld_char1)
    #     page_num_run._r.append(instr_text)
    #     page_num_run._r.append(fld_char2)

    #     of_run = paragraph.add_run()
    #     t2 = self.create_element('w:t')
    #     self.create_attribute(t2, 'xml:space', 'preserve')
    #     t2.text = ' of '
    #     of_run._r.append(t2)

    #     fld_char3 = self.create_element('w:fldChar')
    #     self.create_attribute(fld_char3, 'w:fldCharType', 'begin')

    #     instr_text2 = self.create_element('w:instrText')
    #     self.create_attribute(instr_text2, 'xml:space', 'preserve')
    #     instr_text2.text = "NUMPAGES"

    #     fld_char4 = self.create_element('w:fldChar')
    #     self.create_attribute(fld_char4, 'w:fldCharType', 'end')

    #     num_pages_run = paragraph.add_run()
    #     num_pages_run._r.append(fld_char3)
    #     num_pages_run._r.append(instr_text2)
    #     num_pages_run._r.append(fld_char4)

    # def insert_h_line(self, paragraph):
    #     self.is_not_used()
    #     p = paragraph._p  # p is the <w:p> XML element
    #     p_pr = p.get_or_add_pPr()
    #     p_bdr = OxmlElement('w:pBdr')
    #     p_pr.insert_element_before(p_bdr,
    #                                'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
    #                                'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
    #                                'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
    #                                'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
    #                                'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
    #                                'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
    #                                'w:pPrChange'
    #                                )
    #     bottom = OxmlElement('w:bottom')
    #     bottom.set(qn('w:val'), 'single')
    #     bottom.set(qn('w:sz'), '6')
    #     bottom.set(qn('w:space'), '1')
    #     bottom.set(qn('w:color'), 'auto')
    #     p_bdr.append(bottom)

    # def is_not_used(self):
    #     pass

    # def create_document(self):
    #     self.is_not_used()
    #     doc = Document()
    #     section = doc.sections[-1]
    #     new_width, new_height = section.page_height, section.page_width
    #     section.orientation = WD_ORIENTATION.LANDSCAPE
    #     section.page_width = new_width
    #     section.page_height = new_height

    #     return doc

    # def add_heading(self, text, level):
    #     self.doc.add_heading(text, level)

    # def add_paragraph(self, text, font_size, is_font_bold):
    #     para = self.doc.add_paragraph()
    #     para_run = para.add_run(text)
    #     # Increasing size of the font
    #     para_run.font.size = Pt(font_size)
    #     if is_font_bold:
    #         para_run.bold = True
    #     else:
    #         para_run.bold = False

    #     return para

    # def add_paragraph_run(self, para, text, font_size):
    #     self.is_not_used()
    #     run_para = para.add_run(text)
    #     run_para.font.size = Pt(font_size)
    #     if text == 'Passed':
    #         run_para.font.color.rgb = RGBColor(0, 0xFF, 0)
    #     elif text == 'Failed':
    #         run_para.font.color.rgb = RGBColor(0xFF, 0, 0)

    # def add_table(self, data):
    #     df = pd.DataFrame(data, columns=['Sno', 'Step', 'Expected Result', 'Actual Result', 'Status', 'screenshot'])

    #     t = self.doc.add_table(df.shape[0] + 1, df.shape[1] - 1)
    #     t.style = 'Table Grid'

    #     # add the header rows.
    #     for j in range(df.shape[-1] - 1):
    #         t.cell(0, j).text = df.columns[j]
    #         shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    #         t.cell(0, j)._tc.get_or_add_tcPr().append(shading_elm)

    #     # add the rest of the data frame
    #     for i in range(df.shape[0]):
    #         for j in range(df.shape[-1] - 1):
    #             if j == 3:
    #                 paragraph = t.cell(i + 1, j).paragraphs[0]
    #                 # paragraph = t.cell(i + 1, j).add_paragraph()
    #                 text_run = paragraph.add_run()
    #                 text_run.text = str(df.values[i, j]) + '\n\r\n\r'
    #                 logo_run = paragraph.add_run()
    #                 # logo_run.add_picture("logo.png", width=Inches(5))
    #                 logo_run.add_picture(str(df.values[i, j + 2]), width=Inches(4))
    #                 text_run1 = paragraph.add_run()
    #                 text_run1.text = '\n\r'

    #             elif j == 4:
    #                 paragraph = t.cell(i + 1, j).paragraphs[0]
    #                 text_run = paragraph.add_run()
    #                 text_run.text = str(df.values[i, j])
    #                 if str(df.values[i, j]) == 'Passed':
    #                     text_run.font.color.rgb = RGBColor(0, 0xFF, 0)
    #                 else:
    #                     text_run.font.color.rgb = RGBColor(0xFF, 0, 0)

    #             else:
    #                 t.cell(i + 1, j).text = str(df.values[i, j])

    # def add_footer(self):
    #     self.add_page_number(self.doc.sections[0].footer.paragraphs[0])

    # def add_header(self, document_header):
    #     section = self.doc.sections[0]
    #     header = section.header
    #     paragraph = header.paragraphs[0]
    #     # paragraph.text = "Title of my document"
    #     logo_run = paragraph.add_run()
    #     logo_run.add_picture(self.utils.get_resource_path("logo.png"), width=Inches(1))

    #     text_run = paragraph.add_run()
    #     text_run.text = '\t' + '              ' + document_header  # + '\n\r'
    #     # text_run.text = document_header
    #     text_run.font.size = Pt(40)
    #     text_run.style = "Heading 1 Char"
    #     # text_run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #     self.insert_h_line(paragraph)

    # def save_file(self):
    #     self.doc.save(self.document_name)

    # def convert_to_pdf(self):
    #     self.is_not_used()
    #     wd_format_pdf = 17

    #     in_file = os.path.abspath(self.document_name)
    #     out_file = os.path.abspath(self.document_name_pdf)

    #     word = comtypes.client.CreateObject('Word.Application')
    #     doc_pdf = word.Documents.Open(in_file)
    #     doc_pdf.SaveAs(out_file, FileFormat=wd_format_pdf)
    #     doc_pdf.Close()
    #     word.Quit()

    #     os.remove(in_file)


# doc = Document("your_doc.docx")
# doc = Document()
# section = doc.sections[-1]
# new_width, new_height = section.page_height, section.page_width
# section.orientation = WD_ORIENTATION.LANDSCAPE
# section.page_width = new_width
# section.page_height = new_height

# doc.add_heading('GeeksForGeeks', 0)
#
# # Adding paragraph with Increased font size
# doc.add_heading('Increased Font Size Paragraph:', 3)
# para = doc.add_paragraph().add_run(
#     'GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.')
# # Increasing size of the font
# para.font.size = Pt(12)

# Adding paragraph with normal font size
# doc.add_heading('Normal Font Size Paragraph:', 3)
# doc.add_paragraph(
#     'GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.GeeksforGeeks is a Computer Science portal for geeks.')
# data = [[1, 'Click button', 'The click should be succcessful', 'The click should be succcessful', 'Passed'],
#         [1, 'Click button', 'The click should be succcessful', 'The click should be succcessful', 'Passed'],
#         [1, 'Click button', 'The click should be succcessful', 'The click should be succcessful', 'Passed'],
#         [1, 'Click button', 'The click should be succcessful', 'The click should be succcessful', 'Passed']]
# df = pd.DataFrame(data, columns=['Sno', 'Step', 'Expected Result', 'Actual Result', 'Status'])
#
# t = doc.add_table(df.shape[0] + 1, df.shape[1])
# t.style = 'Table Grid'
#
# # add the header rows.
# for j in range(df.shape[-1]):
#     t.cell(0, j).text = df.columns[j]
#     shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
#     t.cell(0, j)._tc.get_or_add_tcPr().append(shading_elm)
#
# # add the rest of the data frame
# for i in range(df.shape[0]):
#     for j in range(df.shape[-1]):
#         if j == 3:
#             # paragraph = t.cell.paragraphs[0]
#             paragraph = t.cell(i + 1, j).add_paragraph()
#             text_run = paragraph.add_run()
#             text_run.text = "The button is clicked successfully" + '\n\r\n\r'
#             logo_run = paragraph.add_run()
#             logo_run.add_picture("logo.png", width=Inches(5))
#
#         else:
#             t.cell(i + 1, j).text = str(df.values[i, j])

# add_page_number(doc.sections[0].footer.paragraphs[0])
# section = doc.sections[0]
# header = section.header
# paragraph = header.paragraphs[0]
# # paragraph.text = "Title of my document"
# logo_run = paragraph.add_run()
# logo_run.add_picture("logo.png", width=Inches(1))
#
# text_run = paragraph.add_run()
# text_run.text = '\t' + "My Awesome Header" + '\n\r'
# text_run.style = "Heading 1 Char"
# insertHR(paragraph)
# doc.save("your_doc.docx")

# wdFormatPDF = 17
#
# in_file = os.path.abspath("your_doc.docx")
# out_file = os.path.abspath("your_doc.pdf")
#
# word = comtypes.client.CreateObject('Word.Application')
# docpdf = word.Documents.Open(in_file)
# docpdf.SaveAs(out_file, FileFormat=wdFormatPDF)
# docpdf.Close()
# word.Quit()
#
# os.remove(in_file)

# pdf = PdfReporting("test.docx")
# pdf.add_heading("aldldfo aodfdlf dofd fdofdf odfofd ", 0)
# pdf.add_paragraph("thie si fos fp area pf aldof a;;dfoa fdo relf or", 15)
# pdf.add_table(data)
# pdf.add_footer()
# pdf.add_header("Test Results")
# pdf.save_file()
# pdf.convert_to_pdf("test.docx", "test.pdf")
